from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
from datetime import datetime
import os, json, smtplib, logging
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

app = Flask(__name__)
CORS(app)

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.FileHandler("cert_reporting.log"), logging.StreamHandler()]
)
logger = logging.getLogger("cert_reporter")

CONFIG = {
    "template_path": "backend/templates/certin_form.docx",  # Ensure this file exists!
    "output_dir": "/absolute/path/to/filled_reports/",
    "processed_dir": "/absolute/path/to/processed_files/",
    "watch_dir": "/absolute/path/to/incoming_json/",
    "sender_email": "chiragsolanki9821@gmail.com",
    "sender_password": "ekly segp gngi jjzw",  # Use secure methods in production!
    "cert_in_email": "cygnusprofile@gmail.com",
}

# Create directories if they don't exist
for dir_path in [CONFIG["output_dir"], CONFIG["processed_dir"]]:
    os.makedirs(dir_path, exist_ok=True)

# --- Helper functions (unchanged) ---
def get_risk_factors(gemini_analysis):
    if not gemini_analysis or not gemini_analysis.get("risk_factors"):
        return "No specific risk factors identified."
    return "\n".join([f"- {factor}" for factor in gemini_analysis["risk_factors"]])

def format_datetime(timestamp):
    if not timestamp:
        return "Unknown"
    try:
        dt = datetime.fromtimestamp(timestamp)
        return dt.strftime("%d/%m/%Y %H:%M UTC")
    except:
        return "Invalid timestamp"

def extract_location_from_whois(whois_data):
    if not whois_data or not whois_data.get("contacts"):
        return "Unknown"
    contacts = whois_data.get("contacts", {})
    for contact_type in ["owner", "admin", "tech"]:
        if contact_type in contacts and contacts[contact_type]:
            contact = contacts[contact_type][0]
            country = contact.get("country", "")
            state = contact.get("state", "")
            city = contact.get("city", "")
            if country:
                location_parts = [part for part in [city, state, country] if part]
                return ", ".join(location_parts)
    return "Unknown"

def extract_registrar_info(whois_data):
    if not whois_data or not whois_data.get("registrar"):
        return "Unknown"
    registrar = whois_data.get("registrar", {})
    return registrar.get("name", "Unknown")

def extract_ip_info(dns_records):
    if not dns_records or not dns_records.get("A"):
        return "Unknown"
    ip_addresses = dns_records.get("A", [])
    return ", ".join(ip_addresses) if ip_addresses else "Unknown"

def generate_unique_filename(domain, base_dir):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    sanitized_domain = ''.join(c if c.isalnum() else '_' for c in domain)
    return os.path.join(base_dir, f"cert_report_{sanitized_domain}_{timestamp}.docx")

def process_phishing_report(phishing_data):
    try:
        domain = phishing_data.get("domain", "Unknown")
        risk_score = phishing_data.get("risk_score", 0)
        confidence = phishing_data.get("confidence", "Unknown")
        gemini_analysis = phishing_data.get("gemini_analysis", {})
        whois_data = phishing_data.get("whois_data", {})
        dns_records = phishing_data.get("dns_records", {})
        virustotal = phishing_data.get("virustotal", {})
        creation_time = virustotal.get("creation_date", None)
        detection_time = datetime.now().timestamp()
        
        phishing_details = {
            "[PHISHING_URL]": f"http://{domain}",
            "[IP_ADDRESS]": extract_ip_info(dns_records),
            "[OS_DETAILS]": "Unknown (Hosting Server Details Not Available)",
            "[CLOUD_DETAILS]": f"Hosted by {extract_registrar_info(whois_data)}",
            "[AFFECTED_APP]": f"Potential phishing website at {domain}",
            "[LOCATION]": extract_location_from_whois(whois_data),
            "[ISP_DETAILS]": f"{extract_registrar_info(whois_data)}, IP: {extract_ip_info(dns_records)}",
            "[OCCURRENCE_TIME]": format_datetime(creation_time),
            "[DETECTION_TIME]": format_datetime(detection_time),
            "[INCIDENT_DESCRIPTION]": (
                f"A potential phishing website was detected at {domain}. "
                f"Risk Score: {risk_score}/10 (Confidence: {confidence}). "
                f"The site was created on {format_datetime(creation_time)}. "
                f"\n\nRisk Analysis:\n{get_risk_factors(gemini_analysis)}"
                f"\n\nThe domain is registered with {extract_registrar_info(whois_data)}. "
                f"VirusTotal scan shows {virustotal.get('malicious_engines', 0)} engines flagged this domain as malicious "
                f"out of {virustotal.get('total_engines', 0)} total engines. "
                f"Immediate intervention is recommended to prevent potential data theft and financial fraud."
            )
        }
        
        output_path = generate_unique_filename(domain, CONFIG["output_dir"])
        doc = Document(CONFIG["template_path"])
        
        # Replace placeholders in paragraphs and tables
        for para in doc.paragraphs:
            for key, value in phishing_details.items():
                if key in para.text:
                    para.text = para.text.replace(key, value)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in phishing_details.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
                            
        doc.save(output_path)
        logger.info(f"✅ CERT-In report for {domain} has been saved to {output_path}!")
        return phishing_details, output_path
        
    except Exception as e:
        logger.error(f"❌ Error processing phishing report: {e}")
        return None, None

def send_cert_email(sender_email, sender_password, cert_in_email, report_file_path, phishing_details):
    try:
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = cert_in_email
        domain = phishing_details.get('[PHISHING_URL]', 'Unknown URL').replace('http://', '')
        msg['Subject'] = f"URGENT: Phishing Report - {domain}"
        
        email_body = f"""
Dear CERT-In Team,

Please find attached the phishing incident report generated by our automated detection system. This report provides all necessary details for immediate action.

**Key Details:**
- **Phishing URL:** {phishing_details.get('[PHISHING_URL]', 'Unknown')}
- **Risk Score:** {phishing_details.get('[INCIDENT_DESCRIPTION]', '').split('Risk Score: ')[1].split(' ')[0] if 'Risk Score: ' in phishing_details.get('[INCIDENT_DESCRIPTION]', '') else 'Unknown'}
- **IP Address:** {phishing_details.get('[IP_ADDRESS]', 'Unknown')}
- **Incident Detection:** {phishing_details.get('[DETECTION_TIME]', 'Unknown')}

{phishing_details.get('[INCIDENT_DESCRIPTION]', '').split("\n\n")[0]}

We request your urgent intervention to take down this phishing site to protect users from potential fraud.

Best regards,
CatchPhish AI Security Team
contact@catchphish.com
"""
        msg.attach(MIMEText(email_body, 'plain'))
        
        with open(report_file_path, "rb") as attachment:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f"attachment; filename={os.path.basename(report_file_path)}")
            msg.attach(part)
        
        server = smtplib.SMTP("smtp.gmail.com", 587)
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, cert_in_email, msg.as_string())
        server.quit()
        logger.info(f"✅ Phishing report for {domain} successfully sent to {cert_in_email}!")
        return True
    
    except Exception as e:
        logger.error(f"❌ Error sending email: {e}")
        return False

# --- Flask route for CERT-In report generation ---
@app.route('/generate-certin-report', methods=['POST'])
def generate_certin_report():
    try:
        data = request.json
        if 'reportData' not in data:
            return jsonify({'success': False, 'message': 'Missing reportData in request.'}), 400
        
        report_data = json.loads(data['reportData'])
        domain = report_data.get('domain', 'unknown domain')
        
        phishing_details, output_path = process_phishing_report(report_data)
        
        if phishing_details and output_path:
            success = send_cert_email(
                CONFIG["sender_email"], 
                CONFIG["sender_password"], 
                CONFIG["cert_in_email"], 
                output_path, 
                phishing_details
            )
            
            if success:
                return jsonify({
                    'success': True,
                    'message': 'CERT-In report successfully generated and sent.'
                })
            else:
                return jsonify({
                    'success': False,
                    'message': 'Generated report but failed to send email.'
                }), 500
        else:
            return jsonify({
                'success': False,
                'message': 'Failed to process phishing report.'
            }), 500
            
    except Exception as e:
        logger.error(f"Error in generate_certin_report: {e}")
        return jsonify({
            'success': False,
            'message': f'Error processing request: {str(e)}'
        }), 500

if __name__ == '__main__':
    # Ensure you run the Flask app on port 5001 if your front-end is targeting that port
    app.run(host='0.0.0.0', port=5001, debug=True)
