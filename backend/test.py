import requests
import dns.resolver
import json
import re
import os
import numpy as np
from tensorflow import keras
from Feature_Extractor import extract_features
from flask import Flask, request, jsonify, send_from_directory, send_file, render_template_string
from flask_cors import CORS
from bs4 import BeautifulSoup
from google import genai
from docx import Document
import smtplib
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import logging
import io, json, random
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
import sqlite3
import threading
import time
import difflib


VT_API_KEY = "3ea2281a21cf2df9edd36bd5660fa0eaac196e49397ba87d71f118d98982289e"
WHOIS_API_HOST = "zozor54-whois-lookup-v1.p.rapidapi.com"
WHOIS_API_KEY = "d484df19c6msh9303ca1275ac1a8p191fbcjsn68f3d3e05b82"
SCREENSHOT_API_KEY = "0218086a7bca4e76a57da8abb0d5166f"
SCREENSHOT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "screenshots")
MODEL_PATH = "Malicious_URL_Prediction.h5"
GEMINI_API_KEY = "AIzaSyDc4B__rW4_zlwePV5xFiaUDOCBEbHtS0s"
OPENPHISH_FEED_URL = "https://raw.githubusercontent.com/openphish/public_feed/refs/heads/main/feed.txt"

app = Flask(__name__)
CORS(app)

CONFIG = {
    "template_path": "./templates/certin_form.docx",
    "output_dir": "filled_reports/",
    "processed_dir": "processed_files/",
    "watch_dir": "incoming_json/",
    "sender_email": "chiragsolanki9821@gmail.com",
    "sender_password": "ekly segp gngi jjzw",
    "cert_in_email": "cygnusprofile@gmail.com",
}

if not os.path.exists(SCREENSHOT_DIR):
    os.makedirs(SCREENSHOT_DIR)
for dir_path in [CONFIG["output_dir"], CONFIG["processed_dir"]]:
    os.makedirs(dir_path, exist_ok=True)

def get_db_connection():
    conn = sqlite3.connect("reports.db")
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS certin_reports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT,
            domain TEXT,
            report_file_path TEXT,
            phishing_details TEXT,
            status TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    conn.close()

init_db()

def get_dns_records(domain):
    records = {"A": [], "NS": [], "MX": []}
    for record_type in ["A", "NS", "MX"]:
        try:
            records[record_type] = [r.to_text() for r in dns.resolver.resolve(domain, record_type)]
        except:
            records[record_type] = "Not Found"
    return records

def check_virustotal(domain):
    headers = {"x-apikey": VT_API_KEY}
    try:
        response = requests.get(f"https://www.virustotal.com/api/v3/domains/{domain}", headers=headers, timeout=10)
        if response.status_code == 200:
            data = response.json()
            results = data.get("data", {}).get("attributes", {}).get("last_analysis_results", {})
            return {
                "malicious_engines": sum(1 for r in results.values() if r["category"] == "malicious"),
                "total_engines": len(results),
                "creation_date": data.get("data", {}).get("attributes", {}).get("creation_date", "Unknown"),
                "reputation": data.get("data", {}).get("attributes", {}).get("reputation", 0)
            }
    except Exception as e:
        print(f"VirusTotal error: {e}")
    return {"error": "Could not fetch from VirusTotal"}

import urllib.parse

def fetch_openphish():
    try:
        response = requests.get(OPENPHISH_FEED_URL, timeout=30)
        response.raise_for_status()
        urls = response.text.split("\n")
        # Extract domains from URLs
        domains = []
        for url in urls:
            if url:  # Skip empty lines
                try:
                    parsed = urllib.parse.urlparse(url)
                    domain = parsed.netloc
                    if domain:
                        domains.append(domain)
                except:
                    continue
        return domains
    except Exception as e:
        print(f"OpenPhish error: {e}")
        return []

def check_openphish(domain, phishing_list):
    return domain in phishing_list  # Direct domain comparison

def check_suspicious_patterns(domain):
    suspicious_indicators = []
    if domain.count('-') > 0:  # Changed from > 2 to > 0
        suspicious_indicators.append("Contains dashes in domain name")
    if re.search(r"[^a-zA-Z0-9.-]", domain):
        suspicious_indicators.append("Contains unusual symbols")
    return suspicious_indicators

def get_gemini_analysis(url):
    try:
        response = requests.get(url, timeout=10)
        soup = BeautifulSoup(response.content, 'html5lib')
        prompt = f"""
        Analyze the following HTML for security risks:
        - Login form legitimacy
        - Brand impersonation
        - Malicious JavaScript
        - Social engineering
        
        Return JSON format:
        {{
            "risk_score": float,
            "risk_factors": list,
            "confidence": float
        }}
        
        HTML Content:
        {soup.prettify()}
        """
        client = genai.Client(api_key=GEMINI_API_KEY)
        response = client.models.generate_content(model="gemini-2.0-flash", contents=prompt)
        gemini_text = response.text.strip()
        if gemini_text.startswith("```json"):
            gemini_text = gemini_text[7:]
        if gemini_text.endswith("```"):
            gemini_text = gemini_text[:-3]
        gemini_data = json.loads(gemini_text.strip())
        gemini_data = json.loads(json.dumps(gemini_data, default=lambda x: float(x)))
        return gemini_data
    except Exception as e:
        print(f"Error in Gemini analysis: {e}")
        return {"risk_score": 0, "risk_factors": [], "confidence": 0}

def get_ml_prediction(url):
    try:
        model = keras.models.load_model(MODEL_PATH)
        url_features = np.array([extract_features(url)])
        prediction = model.predict(url_features)
        return float(prediction[0][0] * 100)
    except Exception as e:
        print(f"ML error: {e}")
        return 0

def get_whois_data(domain):
    url = f"https://{WHOIS_API_HOST}/"
    params = {"domain": domain, "format": "json", "_forceRefresh": "0"}
    headers = {"x-rapidapi-host": WHOIS_API_HOST, "x-rapidapi-key": WHOIS_API_KEY}
    try:
        response = requests.get(url, headers=headers, params=params, timeout=10)
        return response.json() if response.status_code == 200 else {"error": "WHOIS lookup failed"}
    except Exception as e:
        print(f"WHOIS error: {e}")
        return {"error": "WHOIS lookup error"}

def capture_screenshot(url, filename):
    api_url = "https://api.apiflash.com/v1/urltoimage"
    params = {"access_key": SCREENSHOT_API_KEY, "url": url, "wait_until": "page_loaded"}
    try:
        response = requests.get(api_url, params=params, timeout=30)
        # Check if the response was successful
        if response.status_code == 200:
            filepath = os.path.join(SCREENSHOT_DIR, filename)
            with open(filepath, "wb") as file:
                file.write(response.content)
            return f"/screenshot/{filename}"
        else:
            print(f"Screenshot API error: {response.status_code}")
            return None  # Return None instead of a dict
    except Exception as e:
        print(f"Screenshot error: {e}")
        return None  # Return None instead of a dict

@app.route('/screenshot/<filename>', methods=['GET'])
def get_screenshot(filename):
    # Print for debugging
    print(f"Requested screenshot: {filename}")
    print(f"Looking in directory: {SCREENSHOT_DIR}")
    
    # Check if file exists
    full_path = os.path.join(SCREENSHOT_DIR, filename)
    if not os.path.exists(full_path):
        print(f"File not found: {full_path}")
        return {"error": "File not found"}, 404
    
    # Serve the file
    return send_from_directory(SCREENSHOT_DIR, filename)

@app.route('/analyze', methods=['POST'])
def analyze_url():
    try:
        data = request.get_json()
        if not data or "url" not in data:
            return jsonify({"error": "URL parameter is missing."}), 400
        url = data["url"]
        domain = url.replace("http://", "").replace("https://", "").split("/")[0]
        dns_data = get_dns_records(domain)
        vt_data = check_virustotal(domain)
        phishing_list = fetch_openphish()
        is_phishing_openphish = check_openphish(domain, phishing_list)
        suspicious_indicators = check_suspicious_patterns(domain)
        gemini_analysis = get_gemini_analysis(url)
        ml_prediction = get_ml_prediction(url)
        whois_data = get_whois_data(domain)
        screenshot_filename = f"{domain}.png"
        screenshot_url = capture_screenshot(url, screenshot_filename)
        risk_factors = [min(50, (vt_data.get("malicious_engines", 0) / max(1, vt_data.get("total_engines", 1))) * 100)]
        if is_phishing_openphish:
            risk_factors.append(30)
        risk_factors.append(min(20, len(suspicious_indicators) * 5))
        risk_factors.append(min(30, ml_prediction / 3.33))
        risk_score = sum(risk_factors)
        confidence = "High" if risk_score >= 70 else "Medium" if risk_score >= 40 else "Low"
        return jsonify({
            "domain": domain,
            "dns_records": dns_data,
            "virustotal": vt_data,
            "openphish_flagged": is_phishing_openphish,
            "suspicious_indicators": suspicious_indicators,
            "risk_score": risk_score,
            "confidence": confidence,
            "ml_prediction": ml_prediction,
            "gemini_analysis": gemini_analysis,
            "whois_data": whois_data,
            "screenshot": screenshot_url
        })
    except Exception as e:
        print("Error:", str(e))
        return jsonify({"error": "Internal server error"}), 500

def get_risk_factors(gemini_analysis):
    if not gemini_analysis or not gemini_analysis.get("risk_factors"):
        return "No specific risk factors identified."
    return "\n".join([f"- {factor}" for factor in gemini_analysis["risk_factors"]])

def format_datetime(timestamp):
    if not timestamp:
        return "Unknown"
    try:
        dt = datetime.fromtimestamp(timestamp)
        return dt.strftime("%d/%m/%Y %H:%M UTC")
    except:
        return "Invalid timestamp"

def extract_location_from_whois(whois_data):
    if not whois_data or not whois_data.get("contacts"):
        return "Unknown"
    contacts = whois_data.get("contacts", {})
    for contact_type in ["owner", "admin", "tech"]:
        if contact_type in contacts and contacts[contact_type]:
            contact = contacts[contact_type][0]
            country = contact.get("country", "")
            state = contact.get("state", "")
            city = contact.get("city", "")
            if country:
                location_parts = [part for part in [city, state, country] if part]
                return ", ".join(location_parts)
    return "Unknown"

def extract_registrar_info(whois_data):
    if not whois_data or not whois_data.get("registrar"):
        return "Unknown"
    registrar = whois_data.get("registrar", {})
    return registrar.get("name", "Unknown")

def extract_ip_info(dns_records):
    if not dns_records or not dns_records.get("A"):
        return "Unknown"
    ip_addresses = dns_records.get("A", [])
    return ", ".join(ip_addresses) if ip_addresses else "Unknown"

def generate_unique_filename(domain, base_dir):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    sanitized_domain = ''.join(c if c.isalnum() else '_' for c in domain)
    return os.path.join(base_dir, f"cert_report_{sanitized_domain}_{timestamp}.docx")

def process_phishing_report(phishing_data):
    try:
        domain = phishing_data.get("domain", "Unknown")
        risk_score = phishing_data.get("risk_score", 0)
        confidence = phishing_data.get("confidence", "Unknown")
        gemini_analysis = phishing_data.get("gemini_analysis", {})
        whois_data = phishing_data.get("whois_data", {})
        dns_records = phishing_data.get("dns_records", {})
        virustotal = phishing_data.get("virustotal", {})
        creation_time = virustotal.get("creation_date", None)
        detection_time = datetime.now().timestamp()
        phishing_details = {
            "[PHISHING_URL]": f"http://{domain}",
            "[IP_ADDRESS]": extract_ip_info(dns_records),
            "[OS_DETAILS]": "Unknown (Hosting Server Details Not Available)",
            "[CLOUD_DETAILS]": f"Hosted by {extract_registrar_info(whois_data)}",
            "[AFFECTED_APP]": f"Potential phishing website at {domain}",
            "[LOCATION]": extract_location_from_whois(whois_data),
            "[ISP_DETAILS]": f"{extract_registrar_info(whois_data)}, IP: {extract_ip_info(dns_records)}",
            "[OCCURRENCE_TIME]": format_datetime(creation_time),
            "[DETECTION_TIME]": format_datetime(detection_time),
            "[INCIDENT_DESCRIPTION]": (
                f"A potential phishing website was detected at {domain}. "
                f"Risk Score: {risk_score}/10 (Confidence: {confidence}). "
                f"The site was created on {format_datetime(creation_time)}. "
                f"\n\nRisk Analysis:\n{get_risk_factors(gemini_analysis)}"
                f"\n\nThe domain is registered with {extract_registrar_info(whois_data)}. "
                f"VirusTotal scan shows {virustotal.get('malicious_engines', 0)} engines flagged this domain as malicious "
                f"out of {virustotal.get('total_engines', 0)} total engines. "
                f"Immediate intervention is recommended to prevent potential data theft and financial fraud."
            )
        }
        output_path = generate_unique_filename(domain, CONFIG["output_dir"])
        doc = Document(CONFIG["template_path"])
        for para in doc.paragraphs:
            for key, value in phishing_details.items():
                if key in para.text:
                    para.text = para.text.replace(key, value)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in phishing_details.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
        doc.save(output_path)
        return phishing_details, output_path
    except Exception as e:
        return None, None

def send_cert_email(sender_email, sender_password, cert_in_email, report_file_path, phishing_details):
    try:
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = cert_in_email
        domain = phishing_details.get('[PHISHING_URL]', 'Unknown URL').replace('http://', '')
        msg['Subject'] = f"URGENT: Phishing Report - {domain}"
        email_body = f"""
Dear CERT-In Team,

Please find attached the phishing incident report generated by our automated detection system. This report provides all necessary details for immediate action.

**Key Details:**
- **Phishing URL:** {phishing_details.get('[PHISHING_URL]', 'Unknown')}
- **Risk Score:** {phishing_details.get('[INCIDENT_DESCRIPTION]', '').split('Risk Score: ')[1].split(' ')[0] if 'Risk Score: ' in phishing_details.get('[INCIDENT_DESCRIPTION]', '') else 'Unknown'}
- **IP Address:** {phishing_details.get('[IP_ADDRESS]', 'Unknown')}
- **Incident Detection:** {phishing_details.get('[DETECTION_TIME]', 'Unknown')}

{phishing_details.get('[INCIDENT_DESCRIPTION]', '').split("\n\n")[0]}

We request your urgent intervention to take down this phishing site to protect users from potential fraud.

Best regards,
CatchPhish AI Security Team
contact@catchphish.com
"""
        msg.attach(MIMEText(email_body, 'plain'))
        with open(report_file_path, "rb") as attachment:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f"attachment; filename={os.path.basename(report_file_path)}")
            msg.attach(part)
        server = smtplib.SMTP("smtp.gmail.com", 587)
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, cert_in_email, msg.as_string())
        server.quit()
        return True
    except Exception as e:
        return False

@app.route('/generate-certin-report', methods=['POST'])
def generate_certin_report():
    try:
        data = request.json
        report_data = json.loads(data['reportData'])
        username = data.get("username", "unknown_user")
        domain = report_data.get('domain', 'unknown domain')
        phishing_details, output_path = process_phishing_report(report_data)
        print(phishing_details, output_path)
        if phishing_details and output_path:
            success = send_cert_email(
                CONFIG["sender_email"], 
                CONFIG["sender_password"], 
                CONFIG["cert_in_email"], 
                output_path, 
                phishing_details
            )
            if success:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO certin_reports (username, domain, report_file_path, phishing_details, status)
                    VALUES (?, ?, ?, ?, ?)
                """, (username, domain, output_path, json.dumps(phishing_details), "sent"))
                conn.commit()
                conn.close()
                return jsonify({
                    'success': True,
                    'message': 'CERT-In report successfully generated and sent.'
                })
            else:
                return jsonify({
                    'success': False,
                    'message': 'Generated report but failed to send email.'
                })
        else:
            return jsonify({
                'success': False,
                'message': 'Failed to process phishing report.'
            })
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error processing request: {str(e)}'
        })

@app.route('/get-reports', methods=['GET'])
def get_reports():
    username = request.args.get("username")
    if not username:
        return jsonify({"error": "Username is required."}), 400
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM certin_reports WHERE username = ?", (username,))
        rows = cursor.fetchall()
        conn.close()
        reports = [dict(row) for row in rows]
        return jsonify({"reports": reports})
    except Exception as e:
        return jsonify({"error": "Internal server error"}), 500

def extract_domain(url):
    url = url.replace("http://", "").replace("https://", "")
    return url.split("/")[0]

def is_typosquat(original, candidate):
    if original == candidate:
        return False
    ratio = difflib.SequenceMatcher(None, original, candidate).ratio()
    return 0.8 < ratio < 1.0

def parse_interval(interval_str):
    num = float(re.findall(r"[\d\.]+", interval_str)[0])
    if "hour" in interval_str.lower():
        return int(num * 3600)
    if "minute" in interval_str.lower():
        return int(num * 60)
    return int(num)

def monitor_task(original_domain, interval_seconds):
    reported = set()
    while True:
        feed = fetch_openphish()
        for url in feed:
            if not url.strip():
                continue
            candidate_domain = extract_domain(url)
            if candidate_domain and is_typosquat(original_domain, candidate_domain) and candidate_domain not in reported:
                phishing_data = {
                    "domain": candidate_domain,
                    "dns_records": get_dns_records(candidate_domain),
                    "virustotal": check_virustotal(candidate_domain),
                    "openphish_flagged": True,
                    "suspicious_indicators": check_suspicious_patterns(candidate_domain),
                    "risk_score": 0,
                    "confidence": "Unknown",
                    "ml_prediction": get_ml_prediction(candidate_domain),
                    "gemini_analysis": get_gemini_analysis("http://" + candidate_domain),
                    "whois_data": get_whois_data(candidate_domain),
                    "screenshot": capture_screenshot("http://" + candidate_domain, candidate_domain + ".png")
                }
                phishing_details, output_path = process_phishing_report(phishing_data)
                if phishing_details and output_path:
                    success = send_cert_email(
                        CONFIG["sender_email"],
                        CONFIG["sender_password"],
                        CONFIG["cert_in_email"],
                        output_path,
                        phishing_details
                    )
                    if success:
                        reported.add(candidate_domain)
                        print(f"Reported typosquat candidate: {candidate_domain}")
        time.sleep(interval_seconds)

monitors = {}

@app.route('/monitor-domain', methods=['POST'])
def monitor_domain_endpoint():
    try:
        data = request.get_json()
        if not data or "domain" not in data or "interval" not in data:
            return jsonify({"error": "Both 'domain' and 'interval' parameters are required."}), 400
        domain = data["domain"]
        interval_seconds = parse_interval(data["interval"])
        if domain in monitors:
            return jsonify({"message": f"{domain} is already being monitored."})
        thread = threading.Thread(target=monitor_task, args=(domain, interval_seconds), daemon=True)
        monitors[domain] = thread
        thread.start()
        return jsonify({"message": f"Started monitoring {domain} every {data['interval']}."})
    except Exception as e:
        return jsonify({"error": f"Error starting monitor: {str(e)}"}), 500

class PhishingReportGenerator:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.styles.add(ParagraphStyle(name='CustomTitle', parent=self.styles['Heading1'], fontSize=18, alignment=1, spaceAfter=12))
        self.styles.add(ParagraphStyle(name='Subtitle', parent=self.styles['Heading2'], fontSize=14, textColor=colors.HexColor('#444444'), spaceAfter=12))
        self.styles.add(ParagraphStyle(name='SectionHeading', parent=self.styles['Heading3'], fontSize=12, textColor=colors.HexColor('#246AB3'), spaceBefore=12, spaceAfter=6))
        self.styles.add(ParagraphStyle(name='RiskHigh', parent=self.styles['Normal'], fontSize=11, textColor=colors.HexColor('#CC0000'), fontName='Helvetica-Bold'))
        self.styles.add(ParagraphStyle(name='RiskMedium', parent=self.styles['Normal'], fontSize=11, textColor=colors.HexColor('#FF9900'), fontName='Helvetica-Bold'))
        self.styles.add(ParagraphStyle(name='RiskLow', parent=self.styles['Normal'], fontSize=11, textColor=colors.HexColor('#009900'), fontName='Helvetica-Bold'))
        self.styles.add(ParagraphStyle(name='NormalCustom', parent=self.styles['Normal'], fontSize=10, leading=14))
        self.styles.add(ParagraphStyle(name='NormalSmall', parent=self.styles['Normal'], fontSize=8, leading=10))
        self.styles.add(ParagraphStyle(name='TableHeader', parent=self.styles['Normal'], fontSize=10, fontName='Helvetica-Bold', alignment=1))
        self.styles.add(ParagraphStyle(name='Footer', parent=self.styles['Normal'], fontSize=8, textColor=colors.gray, alignment=1))

    def add_watermark(self, canv, doc):
        canv.saveState()
        try:
            canv.setFillAlpha(0.08)
        except Exception:
            pass
        logo_path = "CatchPhish.png"
        if os.path.exists(logo_path):
            canv.drawImage(logo_path, 100, 200, width=400, height=400, preserveAspectRatio=True)
        try:
            canv.setFillAlpha(1.0)
        except Exception:
            pass
        if os.path.exists(logo_path):
            canv.drawImage(logo_path, 100, 200, width=400, height=400, preserveAspectRatio=True)
        canv.setFont('Helvetica-Bold', 16)
        canv.setFillColor(colors.HexColor('#246AB3'))
        canv.drawString(30, A4[1] - 50, "CatchPhish")
        canv.setFont('Helvetica', 10)
        canv.setFillColor(colors.HexColor('#666666'))
        canv.drawString(30, A4[1] - 65, "Advanced Phishing Detection")
        canv.setStrokeColor(colors.HexColor('#246AB3'))
        canv.setLineWidth(2)
        canv.line(30, A4[1] - 80, A4[0] - 30, A4[1] - 80)
        canv.setFont('Helvetica', 8)
        canv.setFillColor(colors.gray)
        page_num = canv.getPageNumber()
        canv.drawString(A4[0] - 60, 30, f"Page {page_num}")
        canv.drawString(30, 30, f"Report ID: {self.report_id}")
        canv.drawString(30, 20, f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        canv.setStrokeColor(colors.gray)
        canv.setLineWidth(0.5)
        canv.line(30, 45, A4[0] - 30, 45)
        canv.restoreState()

    def get_risk_level_style(self, risk_score):
        try:
            risk_score = float(risk_score)
        except (ValueError, TypeError):
            risk_score = 0.0
        # Adjusted thresholds for risk_score between 0 and 100
        if risk_score >= 70.0:
            return self.styles['RiskHigh']
        elif risk_score >= 40.0:
            return self.styles['RiskMedium']
        else:
            return self.styles['RiskLow']

    def get_risk_level_text(self, risk_score):
        try:
            risk_score = float(risk_score)
        except (ValueError, TypeError):
            risk_score = 0.0
        # Adjusted thresholds for risk_score between 0 and 100
        if risk_score >= 70.0:
            return "HIGH RISK"
        elif risk_score >= 40.0:
            return "MEDIUM RISK"
        else:
            return "LOW RISK"

    def format_datetime(self, timestamp):
        if not timestamp:
            return "Unknown"
        try:
            dt = datetime.fromtimestamp(timestamp)
            return dt.strftime("%d/%m/%Y %H:%M UTC")
        except Exception:
            return "Invalid timestamp"

    def extract_location_from_whois(self, whois_data):
        if not whois_data or not whois_data.get("contacts"):
            return "Unknown"
        contacts = whois_data.get("contacts", {})
        for contact_type in ["owner", "admin", "tech"]:
            if contact_type in contacts and contacts[contact_type]:
                contact = contacts[contact_type][0]
                country = contact.get("country", "")
                state = contact.get("state", "")
                city = contact.get("city", "")
                if country:
                    location_parts = [part for part in [city, state, country] if part]
                    return ", ".join(location_parts)
        return "Unknown"

    def extract_registrar_info(self, whois_data):
        if not whois_data or not whois_data.get("registrar"):
            return "Unknown"
        registrar = whois_data.get("registrar", {})
        return registrar.get("name", "Unknown")

    def extract_ip_info(self, dns_records):
        if not dns_records or not dns_records.get("A"):
            return "Unknown"
        ip_addresses = dns_records.get("A", [])
        return ", ".join(ip_addresses) if ip_addresses else "Unknown"

    def generate_report(self, json_data):
        self.report_id = f"PR-{datetime.now().strftime('%Y%m%d')}-{random.randint(100000, 999999)}"
        data = json_data if isinstance(json_data, dict) else json.loads(json_data)
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(pdf_buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        content = []
        domain = data.get("domain", "Unknown")
        # Ensure risk_score is a float between 0 and 100
        risk_score = float(data.get("risk_score", 0))
        confidence = data.get("confidence", "Unknown")
        ml_prediction = data.get("ml_prediction", 0)
        gemini_analysis = data.get("gemini_analysis", {})
        whois_data = data.get("whois_data", {})
        dns_records = data.get("dns_records", {})
        virustotal = data.get("virustotal", {})

        risk_style = self.get_risk_level_style(risk_score)
        risk_level_text = self.get_risk_level_text(risk_score)

        content.append(Spacer(1, 60))
        content.append(Paragraph("Phishing Detection Report", self.styles["Title"]))
        content.append(Spacer(1, 12))
        content.append(Paragraph(f"Domain Analysis: {domain}", self.styles["Subtitle"]))
        content.append(Paragraph(f"Report Generated: {datetime.now().strftime('%B %d, %Y')}", self.styles["NormalCustom"]))
        content.append(Paragraph(f"Analysis ID: {self.report_id}", self.styles["NormalCustom"]))
        content.append(Spacer(1, 12))
        content.append(Paragraph("<hr/>", self.styles["NormalCustom"]))
        content.append(Paragraph("Executive Summary", self.styles["SectionHeading"]))
        summary_text = (
            f"A detailed security analysis has been conducted on <b>{domain}</b>, which has been "
            f"flagged as a <b>{risk_level_text} domain</b> with a risk score of <b>{risk_score:.2f}/100</b>. "
        )
        if risk_score >= 70.0:
            summary_text += "This domain requires immediate attention as it meets our alert threshold for potential phishing activity."
        elif risk_score >= 40.0:
            summary_text += "This domain should be monitored as it shows some indicators of suspicious activity."
        else:
            summary_text += "This domain shows low risk indicators but should still be monitored for changes."
        content.append(Paragraph(summary_text, self.styles["NormalCustom"]))
        content.append(Spacer(1, 12))
        content.append(Paragraph("<hr/>", self.styles["NormalCustom"]))
        content.append(Paragraph("Risk Assessment", self.styles["SectionHeading"]))
        risk_data = [
            [Paragraph("Metric", self.styles["TableHeader"]),
             Paragraph("Value", self.styles["TableHeader"]),
             Paragraph("Description", self.styles["TableHeader"])],
            [Paragraph("<b>Overall Risk Score</b>", self.styles["NormalCustom"]),
             Paragraph(f"<b>{risk_score:.2f}/100</b>", risk_style),
             Paragraph(f"{risk_level_text} level requiring {'immediate action' if risk_score >= 70.0 else 'monitoring'}", self.styles["NormalCustom"])],
            [Paragraph("<b>Machine Learning Prediction</b>", self.styles["NormalCustom"]),
             Paragraph(f"{ml_prediction:.2f}%", self.styles["NormalCustom"]),
             Paragraph("ML-based probability of phishing behavior", self.styles["NormalCustom"])],
            [Paragraph("<b>Gemini AI Confidence</b>", self.styles["NormalCustom"]),
             Paragraph(f"{confidence} ({gemini_analysis.get('confidence', 0)})", self.styles["NormalCustom"]),
             Paragraph("AI confidence in phishing classification", self.styles["NormalCustom"])],
            [Paragraph("<b>VirusTotal Flags</b>", self.styles["NormalCustom"]),
             Paragraph(f"{virustotal.get('malicious_engines', 0)}/{virustotal.get('total_engines', 0)}", self.styles["NormalCustom"]),
             Paragraph("Number of security engines flagging this domain", self.styles["NormalCustom"])]
        ]
        risk_table = Table(risk_data, colWidths=[doc.width * 0.25, doc.width * 0.25, doc.width * 0.5])
        risk_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E5E5E5')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (1, 0), (1, -1), 'CENTER'),
        ]))
        content.append(risk_table)
        content.append(Spacer(1, 12))
        content.append(Paragraph("<hr/>", self.styles["NormalCustom"]))
        content.append(Paragraph("Domain Information", self.styles["SectionHeading"]))
        registrar = self.extract_registrar_info(whois_data)
        creation_date = whois_data.get("created", "Unknown")
        expiry_date = whois_data.get("expires", "Unknown")
        ip_address = self.extract_ip_info(dns_records)
        location = self.extract_location_from_whois(whois_data)
        privacy_service = "Unknown"
        if (whois_data.get("contacts") and whois_data["contacts"].get("owner") and 
            whois_data["contacts"]["owner"][0].get("organization")):
            privacy_service = whois_data["contacts"]["owner"][0]["organization"]
        domain_data = [
            [Paragraph("Property", self.styles["TableHeader"]), 
             Paragraph("Details", self.styles["TableHeader"])],
            [Paragraph("<b>Domain Name</b>", self.styles["NormalCustom"]),
             Paragraph(domain, self.styles["NormalCustom"])],
            [Paragraph("<b>Creation Date</b>", self.styles["NormalCustom"]),
             Paragraph(creation_date, self.styles["NormalCustom"])],
            [Paragraph("<b>Expiration Date</b>", self.styles["NormalCustom"]),
             Paragraph(expiry_date, self.styles["NormalCustom"])],
            [Paragraph("<b>IP Address</b>", self.styles["NormalCustom"]),
             Paragraph(ip_address, self.styles["NormalCustom"])],
            [Paragraph("<b>Registrar</b>", self.styles["NormalCustom"]),
             Paragraph(registrar, self.styles["NormalCustom"])],
            [Paragraph("<b>WHOIS Privacy</b>", self.styles["NormalCustom"]),
             Paragraph(privacy_service, self.styles["NormalCustom"])],
            [Paragraph("<b>Location</b>", self.styles["NormalCustom"]),
             Paragraph(location, self.styles["NormalCustom"])]
        ]
        domain_table = Table(domain_data, colWidths=[doc.width * 0.3, doc.width * 0.7])
        domain_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E5E5E5')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        content.append(domain_table)
        content.append(Spacer(1, 12))
        content.append(Paragraph("<hr/>", self.styles["NormalCustom"]))
        content.append(Paragraph("AI Analysis Insights", self.styles["SectionHeading"]))
        gemini_score = gemini_analysis.get("score", 0)
        gemini_insights = gemini_analysis.get("insights", [])
        gemini_reason = gemini_analysis.get("reason", "No analysis provided")
        content.append(Paragraph(f"<b>AI Phishing Analysis Score:</b> {gemini_score}/10", self.styles["NormalCustom"]))
        content.append(Paragraph(f"<b>Analysis Summary:</b> {gemini_reason}", self.styles["NormalCustom"]))
        content.append(Spacer(1, 6))
        if gemini_insights:
            content.append(Paragraph("<b>Key Insights:</b>", self.styles["NormalCustom"]))
            for insight in gemini_insights:
                content.append(Paragraph(f"• {insight}", self.styles["NormalCustom"]))
        content.append(Spacer(1, 12))
        content.append(Paragraph("<hr/>", self.styles["NormalCustom"]))
        content.append(Paragraph("Security Scan Results", self.styles["SectionHeading"]))
        vt_positives = virustotal.get("malicious_engines", 0)
        vt_total = virustotal.get("total_engines", 0)
        vt_categories = virustotal.get("categories", {})
        vt_last_analysis = virustotal.get("last_analysis_date", None)
        content.append(Paragraph(f"<b>VirusTotal Detection:</b> {vt_positives}/{vt_total} security vendors flagged this domain as malicious", self.styles["NormalCustom"]))
        if vt_last_analysis:
            formatted_date = self.format_datetime(vt_last_analysis)
            content.append(Paragraph(f"<b>Last Analysis Date:</b> {formatted_date}", self.styles["NormalCustom"]))
        if vt_categories:
            content.append(Paragraph("<b>Domain Categories:</b>", self.styles["NormalCustom"]))
            for vendor, category in vt_categories.items():
                content.append(Paragraph(f"• {vendor}: {category}", self.styles["NormalSmall"]))
        content.append(Spacer(1, 12))
        content.append(Paragraph("<hr/>", self.styles["NormalCustom"]))
        content.append(Paragraph("DNS Records", self.styles["SectionHeading"]))
        if dns_records:
            for record_type, records in dns_records.items():
                if records:
                    content.append(Paragraph(f"<b>{record_type} Records:</b>", self.styles["NormalCustom"]))
                    if isinstance(records, list):
                        for record in records:
                            content.append(Paragraph(f"• {record}", self.styles["NormalSmall"]))
                    else:
                        content.append(Paragraph(f"• {records}", self.styles["NormalSmall"]))
                    content.append(Spacer(1, 4))
        else:
            content.append(Paragraph("No DNS records found for this domain.", self.styles["NormalCustom"]))
        content.append(Spacer(1, 12))
        content.append(Paragraph("<hr/>", self.styles["NormalCustom"]))
        content.append(Paragraph("Recommendations", self.styles["SectionHeading"]))
        # Adjusted recommendation thresholds for risk_score between 0 and 100
        if risk_score >= 70.0:
            recommendations = [
                "Block access to this domain immediately across all network infrastructure",
                "Alert users who may have visited this domain about potential data compromise",
                "Monitor for any credentials that may have been entered on this domain",
                "Add this domain to your organization's blocklist",
                "Consider forensic investigation if any users have interacted with this domain"
            ]
        elif risk_score >= 40.0:
            recommendations = [
                "Implement cautionary blocking of this domain",
                "Monitor any traffic to this domain",
                "Consider warning users about this domain",
                "Add enhanced logging for any interactions with this domain",
                "Re-scan this domain periodically to monitor for changes"
            ]
        else:
            recommendations = [
                "Add this domain to your watchlist for regular monitoring",
                "No immediate action required, but maintain vigilance",
                "Consider periodic rescanning of this domain",
                "Document this domain in your security monitoring system"
            ]
        for recommendation in recommendations:
            content.append(Paragraph(f"• {recommendation}", self.styles["NormalCustom"]))
        content.append(Spacer(1, 24))
        content.append(Paragraph("Disclaimer", self.styles["SectionHeading"]))
        disclaimer_text = (
            "This report was automatically generated based on available data at the time of analysis. "
            "The risk assessment provided is based on algorithmic analysis and should be considered as advisory only. "
            "Security professionals should conduct additional investigation before making final determinations. "
            "CatchPhish Security is not liable for actions taken based on this report."
        )
        content.append(Paragraph(disclaimer_text, self.styles["NormalSmall"]))
        doc.build(content, onFirstPage=self.add_watermark, onLaterPages=self.add_watermark)
        pdf_buffer.seek(0)
        return pdf_buffer
    
@app.route("/generate-pdf", methods=["GET", "POST"])
def generate_pdf():
    if request.method == "POST":
        # Check if the request contains JSON data
        if request.is_json:
            json_data = request.get_json()
        # Fallback to file upload if JSON not provided
        elif "json_file" in request.files:
            file = request.files["json_file"]
            if file.filename == "":
                return "No selected file", 400
            try:
                json_data = json.load(file)
            except Exception as e:
                return f"Invalid JSON: {str(e)}", 400
        else:
            return "No JSON data provided", 400

        generator = PhishingReportGenerator()
        pdf_buffer = generator.generate_report(json_data)
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name="phishing_report.pdf",
            mimetype="application/pdf"
        )
    return render_template_string('''
    <!doctype html>
    <title>Upload JSON for Phishing Report</title>
    <h1>Upload JSON File</h1>
    <form method="post" enctype="multipart/form-data">
      <input type="file" name="json_file">
      <input type="submit" value="Upload">
    </form>
    ''')

from fuzzywuzzy import process
from rapidfuzz import process, fuzz
import threading


# Load domain names from file
LOCAL_DOMAINS = []

def load_local_domains():
    global LOCAL_DOMAINS
    try:
        with open("domain-names.txt", "r") as file:
            LOCAL_DOMAINS = [line.strip().lower() for line in file if line.strip()]
        print("Local domains loaded successfully.")
    except Exception as e:
        print(f"Error loading domains from file: {e}")
        LOCAL_DOMAINS = []

# Preload local domains asynchronously in a background thread
threading.Thread(target=load_local_domains, daemon=True).start()

def load_domains():
    domains = LOCAL_DOMAINS.copy()  # Use preloaded local domains
    # Load domains from the remote URL and preprocess them
    remote_url = "https://raw.githubusercontent.com/Phishing-Database/Phishing.Database/refs/heads/master/phishing-links-NEW-today.txt"
    try:
        response = requests.get(remote_url, timeout=10)
        response.raise_for_status()
        remote_domains = [line.strip().lower() for line in response.text.splitlines() if line.strip()]
        domains.extend(remote_domains)
    except Exception as e:
        print(f"Error loading domains from URL: {e}")
    return domains

def clean_domain(url):
    # Remove http:// or https:// and any path after the domain
    url = re.sub(r"https?://", "", url)
    url = url.split("/")[0]
    return url.lower()

@app.route("/fuzzy-search", methods=["POST"])
def fuzzy_search():
    try:
        data = request.get_json()
        if not data or "domain" not in data:
            return jsonify({"error": "Domain parameter is missing."}), 400
        
        query_domain = clean_domain(data["domain"])
        registered_domains = load_domains_from_api(query_domain)
        
        if not registered_domains:
            return jsonify({"error": "Domain list is empty or could not be loaded."}), 500
        
        # Use RapidFuzz's process.extract for faster fuzzy matching
        # This already returns top matches sorted by score (highest first)
        matches = process.extract(query_domain, registered_domains, scorer=fuzz.ratio, limit=10)
        
        # Return the top 10 matches directly
        return jsonify({"query": query_domain, "matches": matches})
    
    except Exception as e:
        print(f"Error in fuzzy search: {e}")
        return jsonify({"error": "Internal server error"}), 500

def load_domains_from_api(search_term):
    try:
        url = "https://brand-alert.whoisxmlapi.com/api/v2"
        payload = {
            "apiKey": "at_pxagA5NDxSiauY63yBWojnhzfsN19",
            "sinceDate": "2025-03-10",
            "mode": "purchase",
            "withTypos": False,
            "responseFormat": "json",
            "punycode": True,
            "includeSearchTerms": [search_term],
            "excludeSearchTerms": []
        }
        
        response = requests.post(url, json=payload)
        if response.status_code == 200:
            result = response.json()
            # Extract just the domain names from the response
            domains = [domain["domainName"] for domain in result.get("domainsList", [])]
            return domains
        else:
            print(f"API request failed with status code: {response.status_code}")
            return []
    except Exception as e:
        print(f"Error loading domains from API: {e}")
        return []
    
from logo import phishing_api

# Register the phishing detection blueprint
app.register_blueprint(phishing_api, url_prefix='/api/phishing')
    
if __name__ == '__main__':
    app.run(debug=True, port=5001)

